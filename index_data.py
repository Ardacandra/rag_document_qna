"""Index creation helper for the RAG pipeline.

This script loads configuration from environment variables and `configs/default.yaml`.
It reads files from `DATA_DIR`, chunks them, embeds, and writes vectors into a ChromaDB
collection using LlamaIndex's `VectorStoreIndex`.
"""

import os
from pathlib import Path
import sys
import yaml

from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, StorageContext
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Document
from llama_index.core import Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
import chromadb

# Local project paths
BASE_DIR = Path(__file__).resolve().parent
CONFIG_PATH = BASE_DIR / "configs" / "default.yaml"

# Load YAML config if present
config = {}
if CONFIG_PATH.exists():
    try:
        with open(CONFIG_PATH, "r", encoding="utf-8") as fh:
            config = yaml.safe_load(fh) or {}
    except Exception as e:
        print(f"Warning: failed to read config {CONFIG_PATH}: {e}")

# Helper to get config from env or yaml
def cfg(key, default=None):
    return os.getenv(key, config.get(key, default))

# Runtime settings (env overrides YAML)
DATA_DIR = cfg("DATA_DIR", "knowledge_base")
CHROMA_PATH = cfg("CHROMA_PATH", "chroma_db")
COLLECTION_NAME = cfg("COLLECTION_NAME", "advanced_rag_collection")
EMBED_MODEL = cfg("EMBED_MODEL", "BAAI/bge-small-en-v1.5")
CHUNK_SIZE = int(cfg("CHUNK_SIZE", 512))
CHUNK_OVERLAP = int(cfg("CHUNK_OVERLAP", 20))

def create_index():
    """Load documents, create a vector store index, and persist to ChromaDB.

    Exits gracefully with helpful messages on failure.
    """
    data_path = Path(DATA_DIR)
    print(f"Starting indexing process from directory: {data_path}")

    if not data_path.exists():
        print(f"Error: data directory '{data_path}' does not exist. Create it and add documents.")
        return

    # 1. LOAD DOCUMENTS
    # We'll construct Document objects ourselves so we can robustly extract text from PDFs
    docs = []

    # PDF text extractor: prefer pdfplumber, fall back to PyMuPDF (fitz)
    try:
        import pdfplumber
    except Exception:
        pdfplumber = None

    try:
        import fitz  # PyMuPDF
    except Exception:
        fitz = None

    # Optional DOCX reader
    try:
        import docx
    except Exception:
        docx = None

    def extract_text_from_pdf(path: Path) -> str:
        if pdfplumber is not None:
            try:
                with pdfplumber.open(path) as pdf:
                    pages = [p.extract_text() or "" for p in pdf.pages]
                    return "\n".join(pages)
            except Exception:
                # fall through to other methods
                pass
        if fitz is not None:
            try:
                doc = fitz.open(path)
                texts = []
                for page in doc:
                    texts.append(page.get_text())
                return "\n".join(texts)
            except Exception:
                pass
        # Last resort: return empty string
        return ""

    for path in sorted(data_path.glob("**/*")):
        if path.is_dir():
            continue
        suffix = path.suffix.lower()
        text = ""
        if suffix == ".pdf":
            text = extract_text_from_pdf(path)
            if not text.strip():
                print(f"Warning: extracted empty text from PDF {path}; it may be scanned/encrypted.")
        elif suffix in {".txt", ".md"}:
            try:
                text = path.read_text(encoding="utf-8", errors="ignore")
            except Exception as e:
                print(f"Failed reading {path}: {e}")
                continue
        elif suffix == ".docx" and docx is not None:
            try:
                doc = docx.Document(str(path))
                paragraphs = [p.text for p in doc.paragraphs]
                text = "\n".join(paragraphs)
            except Exception as e:
                print(f"Failed reading DOCX {path}: {e}")
                continue
        else:
            # try reading as text fallback
            try:
                text = path.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                print(f"Skipping unsupported binary file: {path}")
                continue

        if not text.strip():
            continue

        # Build a Document object if available, otherwise use a simple dict
        if Document is not None:
            try:
                doc_obj = Document(text=text, doc_id=str(path), metadata={"file_name": path.name})
            except TypeError:
                # older/newer signatures may vary
                try:
                    doc_obj = Document(text=text)
                except Exception:
                    doc_obj = {"text": text, "doc_id": str(path)}
        else:
            doc_obj = {"text": text, "doc_id": str(path)}

        docs.append(doc_obj)

    if not docs:
        print(f"No documents found under {data_path}. Nothing to index.")
        return

    print(f"Loaded {len(docs)} documents.")

    # 2. CONFIGURE COMPONENTS
    print(f"Using embedding model: {EMBED_MODEL} (type={cfg('EMBED_MODEL_TYPE', 'local')})")
    embed_model = HuggingFaceEmbedding(model_name=EMBED_MODEL)
    # Force LlamaIndex to use the local embedding model instead of any remote provider
    try:
        Settings.embed_model = embed_model
    except Exception:
        # If Settings can't be set for any reason, continue — we still pass the embedder explicitly
        pass
    node_parser = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    # 3. CHROMADB SETUP
    if chromadb is None:
        print("Error: 'chromadb' is not installed. Install it with 'pip install chromadb' to persist index.")
        return

    chroma_db_path = Path(CHROMA_PATH)
    chroma_db_path.mkdir(parents=True, exist_ok=True)

    try:
        client = chromadb.PersistentClient(path=str(chroma_db_path))
    except Exception as e:
        print(f"Error initializing chromadb PersistentClient at {chroma_db_path}: {e}")
        return

    try:
        chroma_collection = client.get_or_create_collection(COLLECTION_NAME)
    except Exception as e:
        print(f"Error creating/getting chroma collection '{COLLECTION_NAME}': {e}")
        return

    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    # 4. CREATE AND PERSIST INDEX
    try:
        index = VectorStoreIndex.from_documents(
            docs,
            storage_context=storage_context,
            transformations=[node_parser, embed_model],
        )
    except Exception as e:
        print(f"Error creating/persisting index: {e}")
        return

    print("\n--- Indexing Complete ---")
    print(f"Index successfully saved to: {chroma_db_path}")


if __name__ == "__main__":
    create_index()